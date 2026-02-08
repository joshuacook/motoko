"""Convert markdown text to a DOCX document."""

import io
import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(markdown: str, title: str | None = None) -> bytes:
    """Convert markdown text to DOCX bytes.

    Handles headings, paragraphs, bold, italic, bullet lists,
    numbered lists, code blocks, and horizontal rules.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            _add_code_block(doc, "\n".join(code_lines))
            continue

        # Horizontal rule
        if re.match(r"^(\*{3,}|-{3,}|_{3,})\s*$", line.strip()):
            _add_horizontal_rule(doc)
            i += 1
            continue

        # Heading
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            text = heading_match.group(2).strip()
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Bullet list
        bullet_match = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        if bullet_match:
            text = bullet_match.group(2)
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatting(p, text)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_inline_formatting(p, line)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_inline_formatting(paragraph, text: str):
    """Parse inline markdown (bold, italic, code) and add runs to paragraph."""
    # Pattern matches: **bold**, *italic*, `code`, or plain text
    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*)"  # bold+italic
        r"|(\*\*(.+?)\*\*)"     # bold
        r"|(\*(.+?)\*)"         # italic
        r"|(`(.+?)`)"           # inline code
        r"|([^*`]+)"            # plain text
    )

    for match in pattern.finditer(text):
        if match.group(2):  # bold+italic
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(4):  # bold
            run = paragraph.add_run(match.group(4))
            run.bold = True
        elif match.group(6):  # italic
            run = paragraph.add_run(match.group(6))
            run.italic = True
        elif match.group(8):  # inline code
            run = paragraph.add_run(match.group(8))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif match.group(9):  # plain text
            paragraph.add_run(match.group(9))


def _add_code_block(doc: Document, code: str):
    """Add a code block as a monospace paragraph with gray background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_horizontal_rule(doc: Document):
    """Add a horizontal rule as a thin paragraph border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    # Add a simple text-based separator
    run = p.add_run("─" * 50)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
